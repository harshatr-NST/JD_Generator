import streamlit as st
import pdfplumber
import docx
import pytesseract
from PIL import Image
import re
import io
import json

from transformers import AutoTokenizer, AutoModelForSeq2SeqLM

from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

# =====================================================
# STREAMLIT CONFIG
# =====================================================
st.set_page_config(page_title="JD Generator", layout="wide")

# =====================================================
# TEMPLATE DEFINITIONS
# =====================================================
TEMPLATE_FIELDS = [
    "Company Name",
    "Official Website",
    "Preferred Education",
    "Desired Experience",
    "Designation",
    "Stipend/month (part-time)",
    "Stipend/month",
    "Internship Duration",
    "Roles & Responsibilities",
    "Skills",
    "Joining Location",
    "Joining Month",
    "No. of openings",
    "Selection Process",
]

FIELD_KEYS = {
    "Company Name": "company_name",
    "Official Website": "official_website",
    "Preferred Education": "preferred_education",
    "Desired Experience": "desired_experience",
    "Designation": "designation",
    "Stipend/month (part-time)": "stipend_part_time",
    "Stipend/month": "stipend",
    "Internship Duration": "internship_duration",
    "Roles & Responsibilities": "roles_responsibilities",
    "Skills": "skills",
    "Joining Location": "joining_location",
    "Joining Month": "joining_month",
    "No. of openings": "openings",
    "Selection Process": "selection_process",
}

EMPTY_SCHEMA = {v: "" for v in FIELD_KEYS.values()}

# =====================================================
# LOAD LOCAL OPEN-SOURCE LLM (COMPLETION ONLY)
# =====================================================
@st.cache_resource
def load_llm():
    tokenizer = AutoTokenizer.from_pretrained("google/flan-t5-base")
    model = AutoModelForSeq2SeqLM.from_pretrained("google/flan-t5-base")
    return tokenizer, model

tokenizer, llm_model = load_llm()

# =====================================================
# OCR + TEXT EXTRACTION
# =====================================================
def extract_text(file):
    text = ""

    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text()
                if extracted and len(extracted.strip()) > 100:
                    text += extracted + "\n"
                else:
                    image = page.to_image(resolution=300).original
                    ocr_text = pytesseract.image_to_string(
                        image, config="--psm 6"
                    )
                    text += ocr_text + "\n"

    elif file.type.endswith("wordprocessingml.document"):
        doc = docx.Document(file)
        text = "\n".join(p.text for p in doc.paragraphs)

    elif file.type == "text/plain":
        text = file.read().decode("utf-8")

    return re.sub(r"\n{2,}", "\n", text).strip()

# =====================================================
# RULE-BASED EXTRACTION (PRIMARY)
# =====================================================
def rule_extract(text):
    data = EMPTY_SCHEMA.copy()
    sections = {}
    current = None

    for line in text.splitlines():
        l = line.strip()
        if not l:
            continue

        if re.search(r"(role|responsibilit)", l, re.I):
            current = "roles_responsibilities"
            sections[current] = ""
        elif re.search(r"(skill|qualification|requirement)", l, re.I):
            current = "skills"
            sections[current] = ""
        elif re.search(r"(selection|interview|process)", l, re.I):
            current = "selection_process"
            sections[current] = ""
        elif current:
            sections[current] += l + "\n"

    data.update(sections)

    if m := re.search(r"(job title|designation|role)\s*[:\-]\s*(.*)", text, re.I):
        data["designation"] = m.group(2).strip()

    if m := re.search(r"\d+\+?\s*years?", text, re.I):
        data["desired_experience"] = m.group(0)

    if m := re.search(r"(₹|\$)\s?\d+[,\d]*", text):
        data["stipend"] = m.group(0)

    if m := re.search(r"\d+\s*(months?|weeks?)", text, re.I):
        data["internship_duration"] = m.group(0)

    if m := re.search(r"\d+\s+(openings|positions|vacancies)", text, re.I):
        data["openings"] = m.group(0)

    if m := re.search(r"(location|based at)\s*[:\-]?\s*(.*)", text, re.I):
        data["joining_location"] = m.group(2)

    return data

# =====================================================
# SKILL NORMALIZATION
# =====================================================
SKILL_MAP = {
    "Python": ["python", "py"],
    "Java": ["java", "java se", "java 8"],
    "SQL": ["sql", "mysql", "postgres"],
    "Machine Learning": ["machine learning", "ml"],
    "Data Analysis": ["data analysis", "analytics"],
}

def normalize_skills(text):
    found = set()
    lower = text.lower()
    for canon, variants in SKILL_MAP.items():
        for v in variants:
            if v in lower:
                found.add(canon)
    return ", ".join(sorted(found))

# =====================================================
# LLM FILL ONLY BLANK FIELDS
# =====================================================
def llm_fill_missing(raw_text, data):
    missing = [k for k, v in data.items() if not v.strip()]
    if not missing:
        return data

    prompt = f"""
Fill ONLY the missing fields below.
Do not modify existing values.
Return ONLY valid JSON.

Missing fields:
{missing}

Job Description:
{raw_text}

Current data:
{json.dumps(data)}
"""

    inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=2048)
    outputs = llm_model.generate(**inputs, max_new_tokens=512)
    decoded = tokenizer.decode(outputs[0], skip_special_tokens=True)

    try:
        filled = json.loads(decoded)
        for k in missing:
            if k in filled and filled[k]:
                data[k] = filled[k]
    except:
        pass

    return data

# =====================================================
# CONFIDENCE SCORING
# =====================================================
def confidence_score(value):
    if not value:
        return 0.0
    if len(value.split()) > 20:
        return 0.9
    if len(value.split()) > 5:
        return 0.7
    return 0.5

# =====================================================
# PDF GENERATION
# =====================================================
def generate_pdf(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)

    left = ParagraphStyle("left", fontSize=8, textColor=colors.white)
    right = ParagraphStyle("right", fontSize=8)

    table_data = []

    for label in TEMPLATE_FIELDS:
        key = FIELD_KEYS[label]
        table_data.append([
            Paragraph(label, left),
            Paragraph(data.get(key, "").replace("\n", "<br/>"), right)
        ])

    table = Table(table_data, colWidths=[170, 340])
    table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (0,-1), colors.HexColor("#2e74b5")),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("PADDING", (0,0), (-1,-1), 6),
    ]))

    doc.build([table])
    buffer.seek(0)
    return buffer

# =====================================================
# DOCX GENERATION
# =====================================================
def generate_docx(data):
    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = docx.shared.Pt(10)

    for label in TEMPLATE_FIELDS:
        key = FIELD_KEYS[label]
        p = document.add_paragraph()
        run = p.add_run(f"{label}:\n")
        run.bold = True
        document.add_paragraph(data.get(key, ""))

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf

# =====================================================
# STREAMLIT UI
# =====================================================
st.title("JD Generator (OCR + AI Assisted)")

uploaded_file = st.file_uploader(
    "Upload Job Description (PDF / DOCX / TXT)",
    type=["pdf", "docx", "txt"]
)

if uploaded_file:
    raw_text = extract_text(uploaded_file)

    st.subheader("OCR / Extracted Text")
    st.text_area("Extracted Content", raw_text, height=250)

    if st.button("Extract & Auto-Fill"):
        base = rule_extract(raw_text)
        final = llm_fill_missing(raw_text, base)

        if final.get("skills"):
            final["skills"] = normalize_skills(final["skills"])

        st.session_state["data"] = final
        st.session_state["confidence"] = {
            k: confidence_score(v) for k, v in final.items()
        }

if "data" in st.session_state:
    st.subheader("Review & Edit")

    edited = {}
    for label in TEMPLATE_FIELDS:
        key = FIELD_KEYS[label]
        if label in ["Roles & Responsibilities", "Skills", "Selection Process"]:
            edited[key] = st.text_area(label, st.session_state["data"].get(key, ""), height=120)
        else:
            edited[key] = st.text_input(label, st.session_state["data"].get(key, ""))

        st.caption(f"Confidence: {int(st.session_state['confidence'][key]*100)}%")

    pdf = generate_pdf(edited)
    docx_file = generate_docx(edited)

    company = edited.get("company_name", "Company").replace(" ", "_")
    role = edited.get("designation", "Role").replace(" ", "_")

    st.download_button("Download PDF", pdf, f"{company}-{role}.pdf", "application/pdf")
    st.download_button(
        "Download DOCX",
        docx_file,
        f"{company}-{role}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
